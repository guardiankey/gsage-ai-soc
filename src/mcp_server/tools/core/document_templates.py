"""gSage AI — document_templates MCP tool.

Lists document templates available to the current org and downloads individual
templates into the conversation context.

Scope filtering
---------------
- ``organization`` — templates with scope="organization" (visible to all org members)
- ``department``   — templates with scope="department" visible to the current department
- ``user``          — templates uploaded by the current user (scope="user")
- ``all``           — organization + department (when in a dept) + user templates (default)

Templates are GSageFile rows with category="template".
"""

from __future__ import annotations

import time
import logging
from typing import ClassVar, Optional

from sqlalchemy import func, or_, select

from src.mcp_server.tools.base import BaseTool, ToolResult
from src.shared.security.context import AgentContext

log = logging.getLogger(__name__)

_MIME_DOCX = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_MIME_MD = "text/markdown"
_MIME_ZIP = "application/zip"


class DocumentTemplatesTool(BaseTool):
    """
    List and download document templates available to the current organisation.

    Templates are uploaded via the Files API (``POST /v1/orgs/{org_id}/files/upload-template``).
    Each template has a ``scope``:

    - ``organization`` — visible to all members of the org.
    - ``department``    — visible to all members of the current department.
    - ``user``          — private to the uploader.

    Use the returned ``template_id`` with ``generate_document`` to produce
    a document from a template.

    **Built-in templates**

    In addition to user-uploaded templates, ``generate_document`` also
    accepts the following packaged built-ins (not listed by this tool):

    - ``builtin:default`` — minimal Markdown template with gSage CSS
      (works for ``md``/``html``/``docx``/``pdf`` output).
    - ``builtin:pandoc_gsage`` — Pandoc/LaTeX bundle with cover page, TOC
      and gSage colors (``pdf`` output only).

    When ``template_id`` is omitted in ``generate_document``, a built-in
    is auto-selected based on ``output_format`` and ``pandoc``. For CSV
    output no template is needed at all.

    **Downloading a template into the conversation**

    Pass ``download_template_id`` with a ``template_id`` obtained from a
    previous listing to download the template in its original format into
    the conversation context.  The tool creates an artifact (file reference)
    and returns its metadata + ``file_id``.  Use ``read_file`` on the
    returned ``file.file_id`` to inspect the template structure and
    placeholders — the content is NOT returned inline to avoid consuming
    context tokens.

    Optional parameters
    -------------------
    scope (str):
        Filter by template scope.
        ``"all"`` (default), ``"organization"``, ``"department"``, or ``"user"``.
        Ignored when ``download_template_id`` is set.
    content_type (str):
        Filter by MIME type (e.g. ``"application/vnd.openxmlformats-officedocument.wordprocessingml.document"``
        for DOCX, ``"text/markdown"`` for Markdown templates, or
        ``"application/zip"`` for Pandoc bundles).
        Ignored when ``download_template_id`` is set.
    download_template_id (str):
        When set, download a single template by its ``template_id`` instead of
        listing all templates.  The template is stored as a conversation
        artifact — use ``read_file`` on the returned ``file.file_id`` to
        inspect its content.  Template content is NOT returned inline.

    Permission: ``files:read`` (list), ``files:read`` + ``files:write`` (download)
    """

    name: ClassVar[str] = "document_templates"
    version: ClassVar[str] = "1.3.0"
    summary: ClassVar[str] = (
        "List (with optional search filter) and download document templates "
        "(DOCX, Markdown) available to the organization for use with generate_document"
    )
    category: ClassVar[str] = "document"
    permissions: ClassVar[list[str]] = ["files:read", "files:write"]
    rate_limit_per_minute: ClassVar[int] = 60
    timeout_seconds: ClassVar[int] = 15
    use_circuit_breaker: ClassVar[bool] = False
    always_background: ClassVar[bool] = False

    params_schema: ClassVar[dict] = {
        "type": "object",
        "required": [],
        "properties": {
            "scope": {
                "type": "string",
                "enum": ["all", "organization", "department", "user"],
                "description": (
                    "Which templates to list. "
                    "'all' (default): org-wide + department + personal templates. "
                    "'organization': only org-wide templates. "
                    "'department': only templates shared with the current department. "
                    "'user': only templates uploaded by the current user. "
                    "Ignored when 'download_template_id' is set."
                ),
            },
            "content_type": {
                "type": "string",
                "description": (
                    "Optional MIME type filter. "
                    "Example: 'text/markdown' for Markdown templates, "
                    "'application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document' for DOCX templates. "
                    "Ignored when 'download_template_id' is set."
                ),
            },
            "include_variables": {
                "type": "boolean",
                "default": True,
                "description": (
                    "When true (default), download each template and extract the "
                    "placeholder variable names (e.g. '{{title}}', '{{content}}'). "
                    "These variable names MUST be passed via the 'variables' parameter "
                    "when calling generate_document. "
                    "Set to false to skip variable extraction for faster listing. "
                    "Ignored when 'download_template_id' is set (variables are "
                    "always extracted for the downloaded template)."
                ),
            },
            "download_template_id": {
                "type": "string",
                "description": (
                    "When set, download a single template by its UUID (obtained "
                    "from a previous listing) into the conversation context. "
                    "The template is stored as a conversation artifact and its "
                    "metadata + file reference are returned. "
                    "Use read_file on the returned file.file_id to inspect "
                    "the template content — it is NOT returned inline. "
                    "When this parameter is provided, 'scope', 'content_type', "
                    "'search', and 'include_variables' are ignored."
                ),
            },
            "search": {
                "type": "string",
                "description": (
                    "Optional case-insensitive substring filter on template "
                    "filename and description. "
                    "Example: 'ata' matches 'exemplo_ata_assefaz.md' or a "
                    "template with 'Modelo de ata' in its description. "
                    "Ignored when 'download_template_id' is set."
                ),
            },
        },
        "additionalProperties": False,
    }

    async def execute(
        self,
        agent_context: AgentContext,
        params: dict,
        config: dict,
        state: dict,
    ) -> ToolResult:
        from src.shared.database import _get_session_maker  # noqa: PLC0415

        t0 = time.monotonic()

        # ── Download mode ────────────────────────────────────────────────
        download_id: str = str(params.get("download_template_id") or "").strip()
        if download_id:
            return await self._download_single_template(
                template_id=download_id,
                agent_context=agent_context,
                t0=t0,
                session_maker_factory=_get_session_maker,
            )

        # ── List mode ────────────────────────────────────────────────────
        from src.shared.models.generated_file import GSageFile  # noqa: PLC0415

        scope: str = params.get("scope", "all")
        content_type_filter: Optional[str] = params.get("content_type")
        include_variables: bool = bool(params.get("include_variables", True))
        search: str = (params.get("search") or "").strip()

        if scope not in ("all", "organization", "department", "user"):
            scope = "all"

        async with _get_session_maker()() as db:
            stmt = select(GSageFile).where(
                GSageFile.org_id == agent_context.org_id,
                GSageFile.category == "template",
                GSageFile.purged_at.is_(None),
            )

            if scope == "organization":
                stmt = stmt.where(GSageFile.scope == "organization")
            elif scope == "department":
                if agent_context.dept_id is None:
                    return self._failure(
                        "MISSING_CONTEXT",
                        "Department scope requires a department context. The current session has no department assigned.",
                        execution_time_ms=int((time.monotonic() - t0) * 1000),
                    )
                stmt = stmt.where(
                    GSageFile.scope == "department",
                    GSageFile.dept_id == agent_context.dept_id,
                )
            elif scope == "user":
                stmt = stmt.where(
                    GSageFile.scope == "user",
                    GSageFile.user_id == agent_context.user_id,
                )
            else:
                # "all": org templates + dept templates (when in a dept) + user's own
                dept_conditions = [
                    GSageFile.scope == "organization",
                    (
                        GSageFile.scope == "user"
                    ) & (GSageFile.user_id == agent_context.user_id),
                ]
                if agent_context.dept_id is not None:
                    dept_conditions.append(
                        (GSageFile.scope == "department")
                        & (GSageFile.dept_id == agent_context.dept_id)
                    )
                stmt = stmt.where(or_(*dept_conditions))

            if content_type_filter:
                stmt = stmt.where(GSageFile.content_type == content_type_filter)

            if search:
                pattern = f"%{search}%"
                stmt = stmt.where(
                    or_(
                        GSageFile.filename.ilike(pattern),
                        GSageFile.description.ilike(pattern),
                    )
                )

            stmt = stmt.order_by(GSageFile.created_at.desc())
            result = await db.execute(stmt)
            rows = result.scalars().all()

        templates: list[dict] = []
        for row in rows:
            entry: dict = {
                "template_id": str(row.id),
                "filename": row.filename,
                "description": row.description,
                "content_type": row.content_type,
                "scope": row.scope,
                "size_bytes": row.size_bytes,
                "created_at": row.created_at.isoformat() if row.created_at else None,
            }

            if include_variables:
                variables = await self._extract_variables(
                    file_id=str(row.id),
                    org_id=str(agent_context.org_id),
                    content_type=row.content_type,
                    user_id=str(agent_context.user_id),
                    dept_id=str(agent_context.dept_id) if agent_context.dept_id else None,
                )
                entry["variables"] = variables

            templates.append(entry)

        return self._success(
            data={"templates": templates, "count": len(templates)},
            execution_time_ms=int((time.monotonic() - t0) * 1000),
        )

    # ------------------------------------------------------------------
    # Download single template
    # ------------------------------------------------------------------

    async def _download_single_template(
        self,
        template_id: str,
        agent_context: AgentContext,
        t0: float,
        session_maker_factory,
    ) -> ToolResult:
        """Download a single template into the conversation context.

        Loads the template bytes via ``_load_file``, stores a copy as a
        conversation artifact via ``_store_file``, and returns the artifact
        reference plus inline text for Markdown templates.
        """
        # 1. Load the template bytes
        load_result = await self._load_file(
            file_id=template_id,
            org_id=str(agent_context.org_id),
            user_id=str(agent_context.user_id),
            dept_id=str(agent_context.dept_id) if agent_context.dept_id else None,
            max_bytes=10 * 1024 * 1024,  # 10 MB — same cap as generate_document
        )
        if load_result is None:
            return self._failure(
                "TEMPLATE_NOT_FOUND",
                f"Template '{template_id}' not found or access denied. "
                "Verify the template_id with a listing call first.",
                execution_time_ms=int((time.monotonic() - t0) * 1000),
            )

        template_bytes: bytes = load_result["data"]
        filename: str = load_result["filename"]
        content_type: str = load_result["content_type"]
        size_bytes: int = load_result["size_bytes"]
        truncated: bool = load_result.get("truncated", False)

        # 2. Extract variables
        variables: list[str] = []
        if content_type == _MIME_MD:
            variables = self._vars_from_md(
                template_bytes.decode("utf-8", errors="replace")
            )
        elif content_type in (_MIME_DOCX, _MIME_ZIP):
            variables = self._vars_from_docx(template_bytes)

        # 3. Store a copy as a conversation artifact
        file_info: Optional[dict] = None
        async with session_maker_factory()() as db_session:
            file_info = await self._store_file(
                data=template_bytes,
                filename=filename,
                content_type=content_type,
                agent_context=agent_context,
                session=db_session,
                description=f"Template: {filename}",
                scope="user",
            )
            if file_info is None:
                return self._failure(
                    "STORE_FAILED",
                    f"Failed to store template '{filename}' as a conversation artifact. "
                    "The template was loaded but could not be saved to the file store.",
                    retryable=True,
                    execution_time_ms=int((time.monotonic() - t0) * 1000),
                )

        # 4. Build response — NEVER include content inline.
        #    The agent must use read_file on file.file_id to inspect
        #    the template.  Returning 60+ KB of Markdown inline would
        #    waste context tokens and duplicate data already available
        #    via the file store.
        data: dict = {
            "template_id": template_id,
            "filename": filename,
            "content_type": content_type,
            "size_bytes": size_bytes,
            "variables": variables,
            "file": file_info,
        }

        return self._success(
            data=data,
            execution_time_ms=int((time.monotonic() - t0) * 1000),
        )

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    async def _extract_variables(
        self,
        file_id: str,
        org_id: str,
        content_type: str,
        user_id: str | None = None,
        dept_id: str | None = None,
    ) -> list[str]:
        """Download a template and extract placeholder variable names."""
        load_result = await self._load_file(
            file_id=file_id,
            org_id=org_id,
            user_id=user_id,
            dept_id=dept_id,
            max_bytes=2 * 1024 * 1024,
        )
        if load_result is None:
            return []

        data: bytes = load_result["data"]

        if content_type == "text/markdown":
            return self._vars_from_md(data.decode("utf-8", errors="replace"))

        if content_type in (_MIME_DOCX, "application/zip"):
            return self._vars_from_docx(data)

        return []

    @staticmethod
    def _vars_from_md(text: str) -> list[str]:
        """Extract Jinja2 variable names from Markdown template text."""
        from src.shared.services.document_converter import extract_template_variables
        return extract_template_variables(text)

    @staticmethod
    def _vars_from_docx(data: bytes) -> list[str]:
        """Extract {{placeholder}} names from a DOCX template."""
        import io as _io
        import re as _re

        try:
            from docx import Document
        except ImportError:
            return []

        try:
            doc = Document(_io.BytesIO(data))
        except Exception:
            return []

        pattern = _re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_.]*)\s*\}\}")
        found: set[str] = set()

        for para in doc.paragraphs:
            full = "".join(run.text for run in para.runs)
            found.update(pattern.findall(full))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full = "".join(run.text for run in para.runs)
                        found.update(pattern.findall(full))

        found.add("content")
        return sorted(found)
