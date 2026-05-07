"""gSage AI — Document conversion utilities.

Stateless helpers that convert between Markdown, HTML, DOCX and PDF.
All functions are pure (no DB, no MinIO) and can be called from tools or tasks.

PDF and DOCX conversion requires pandoc to be installed in the container.
HTML conversion uses the ``markdown`` library (pure Python, no system deps).
Jinja2 template rendering uses SandboxedEnvironment to prevent injection attacks.

Typical pipeline
----------------
Markdown template  →  render_jinja2_template  →  md_to_html  →  html_to_pdf
DOCX template      →  fill_docx_template      →  docx_to_pdf
"""

from __future__ import annotations

import asyncio
import io
import logging
import re
import unicodedata
from typing import Optional

log = logging.getLogger(__name__)

# Front-matter delimiter
_FRONT_MATTER_RE = re.compile(r"^\s*---\s*\n(.*?)\n---\s*\n", re.DOTALL)

# Jinja2 variable reference: {{ var_name }}  (also {{ obj.attr }})
_JINJA2_VAR_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_.]*)\s*\}\}")

# Pandoc binary name
_PANDOC_BIN = "pandoc"

# Subprocess timeout (seconds) for pandoc calls
_PANDOC_TIMEOUT = 30

# Unicode code points in the "Symbol, Other" (So) category that pdflatex
# can typeset via the standard T1 + inputenc utf8 setup without extra packages.
_LATEX_SAFE_SO: frozenset[int] = frozenset([
    0x00A9,  # © COPYRIGHT SIGN
    0x00AE,  # ® REGISTERED SIGN
    0x00B0,  # ° DEGREE SIGN
    0x2122,  # ™ TRADE MARK SIGN  (\texttrademark via textcomp)
])


def _is_latex_safe(ch: str) -> bool:
    """Return True if *ch* can be typeset by pdflatex with inputenc utf8 + fontenc T1.

    Strips:
    - Control characters (except \\t \\n \\r)
    - Non-BMP code points (U+10000+)
    - Unicode category "So" (Symbol, Other): emoji, pictographs, ⏳ ✅ ❌ ☀ ⭐ …
      Exceptions: © ® ° ™ which pdflatex handles via T1/textcomp.
    - Unicode category "Sm" above Latin-1 (U+00FF): → ← ∑ ∫ ∞ …
      (require amssymb/unicode-math not loaded by default)
    - Unicode categories "Cs" / "Co" / "Cn": surrogates, private use, unassigned
    """
    cp = ord(ch)
    # Whitespace explicitly kept
    if cp in (0x09, 0x0A, 0x0D):
        return True
    # Other control chars and DEL
    if cp < 0x20 or cp == 0x7F:
        return False
    # ASCII printable — always safe
    if cp <= 0x7E:
        return True
    # Non-BMP supplementary planes
    if cp > 0xFFFF:
        return False
    cat = unicodedata.category(ch)
    # Surrogates, private use, unassigned
    if cat in ("Cs", "Co", "Cn"):
        return False
    # Symbol, Other: strip unless known to work in pdflatex
    if cat == "So" and cp not in _LATEX_SAFE_SO:
        return False
    # Symbol, Math above Latin-1 Supplement: arrows, ∑, ∫ …
    if cat == "Sm" and cp > 0x00FF:
        return False
    return True


# ---------------------------------------------------------------------------
# Public helpers
# ---------------------------------------------------------------------------


def parse_md_front_matter(md_text: str) -> tuple[dict, str]:
    """Extract YAML front-matter from a Markdown string.

    Returns a tuple of ``(metadata_dict, body_without_front_matter)``.
    If no front-matter block is present, returns ``({}, md_text)``.

    Only scalar values (str, int, float, bool) and simple lists are
    supported — complex nested structures are returned as raw strings.
    """
    match = _FRONT_MATTER_RE.match(md_text)
    if not match:
        return {}, md_text

    raw_yaml = match.group(1)
    body = md_text[match.end():]

    metadata: dict = {}
    for line in raw_yaml.splitlines():
        if ":" not in line:
            continue
        key, _, value = line.partition(":")
        key = key.strip()
        value = value.strip()
        # Strip surrounding quotes
        if (value.startswith('"') and value.endswith('"')) or (
            value.startswith("'") and value.endswith("'")
        ):
            value = value[1:-1]
        metadata[key] = value

    return metadata, body


def extract_template_variables(template_text: str) -> list[str]:
    """Return sorted unique Jinja2 variable names found in *template_text*.

    Only simple ``{{ variable }}`` references are extracted — Jinja2 control
    structures (``{% … %}``) and filters (``{{ x | filter }}``) are ignored.
    The reserved ``content`` variable is always included in the result.
    """
    _, body = parse_md_front_matter(template_text)
    found = set(_JINJA2_VAR_RE.findall(body))
    found.add("content")
    return sorted(found)


def find_latex_unsafe_chars(text: str) -> list[str]:
    """Return sorted list of unique characters in *text* that pdflatex cannot typeset.

    Useful for generating user-facing warnings before PDF generation.
    """
    seen: set[str] = set()
    result: list[str] = []
    for ch in text:
        if ch not in seen and not _is_latex_safe(ch) and ch not in ("\t", "\n", "\r"):
            seen.add(ch)
            result.append(ch)
    return result


def strip_non_bmp(text: str) -> str:
    """Remove characters that cause LaTeX errors when pandoc generates PDF.

    Uses :func:`_is_latex_safe` (Unicode-category based) instead of a fixed
    regex, so newly-assigned emoji / symbol code-points are handled
    automatically without updating a block-list.

    Strips:
    - Non-BMP supplementary-plane characters (U+10000+, e.g. 📊 ⏳)
    - BMP Symbol, Other ("So"): emoji, pictographs, ⏳ ✅ ❌ ☀ ⭐ …
    - BMP Symbol, Math ("Sm") above Latin-1: → ← ∑ ∫ …
    - Surrogates, private use, unassigned characters
    - Control characters (except \\t \\n \\r)
    - Invalid UTF-8 byte sequences
    """
    if isinstance(text, bytes):
        text = text.decode("utf-8", errors="replace")
    else:
        text = text.encode("utf-8", errors="replace").decode("utf-8")

    return "".join(ch for ch in text if _is_latex_safe(ch))


def render_jinja2_template(template_text: str, variables: dict) -> str:
    """Render *template_text* as a Jinja2 template with *variables*.

    Uses ``SandboxedEnvironment`` to prevent arbitrary code execution.

    Parameters
    ----------
    template_text:
        Jinja2 template string (may include ``{{ var }}``, ``{% for … %}``, etc.).
    variables:
        Mapping of variable names to values passed to the template.

    Returns
    -------
    str
        Rendered output.

    Raises
    ------
    jinja2.TemplateError
        If the template is syntactically invalid or rendering fails.
    """
    from jinja2.sandbox import SandboxedEnvironment

    env = SandboxedEnvironment(autoescape=False)
    tmpl = env.from_string(template_text)
    return tmpl.render(**variables)


def md_to_html(md_text: str, css: Optional[str] = None) -> str:
    """Convert Markdown to an HTML string.

    The ``markdown`` library (pure Python) handles conversion.
    Supported extensions: ``tables``, ``fenced_code``, ``toc``.

    Parameters
    ----------
    md_text:
        Input Markdown text.
    css:
        Optional CSS string embedded as a ``<style>`` block in the ``<head>``.

    Returns
    -------
    str
        Complete HTML document (``<html><head>…</head><body>…</body></html>``).
    """
    import markdown as md_lib

    extensions = ["tables", "fenced_code", "toc"]
    body_html = md_lib.markdown(md_text, extensions=extensions)

    style_block = f"<style>\n{css}\n</style>\n" if css else ""
    return (
        "<!DOCTYPE html>\n"
        "<html>\n"
        "<head>\n"
        "<meta charset=\"utf-8\">\n"
        f"{style_block}"
        "</head>\n"
        "<body>\n"
        f"{body_html}\n"
        "</body>\n"
        "</html>\n"
    )


async def html_to_pdf(html: str) -> bytes:
    """Convert an HTML string to PDF bytes using pandoc.

    pandoc reads from stdin and writes PDF to stdout (via ``--output=-``).
    Non-BMP characters (emoji, supplementary symbols) are stripped before
    conversion to avoid LaTeX encoding errors.

    Parameters
    ----------
    html:
        Complete HTML document string.

    Returns
    -------
    bytes
        Raw PDF bytes.

    Raises
    ------
    FileNotFoundError
        If pandoc is not installed.
    RuntimeError
        If pandoc exits with a non-zero status.
    """
    cmd = [_PANDOC_BIN, "--from=html", "--to=pdf", "--output=-"]
    return await _run_pandoc(cmd, input_data=strip_non_bmp(html).encode("utf-8"))


async def html_to_docx(html: str) -> bytes:
    """Convert an HTML string to DOCX bytes using pandoc.

    Non-BMP characters (emoji, supplementary symbols) are stripped before
    conversion for consistency and to avoid potential encoding issues.

    Parameters
    ----------
    html:
        Complete HTML document string.

    Returns
    -------
    bytes
        Raw DOCX bytes (Open XML format).

    Raises
    ------
    FileNotFoundError
        If pandoc is not installed.
    RuntimeError
        If pandoc exits with a non-zero status.
    """
    cmd = [_PANDOC_BIN, "--from=html", "--to=docx", "--output=-"]
    return await _run_pandoc(cmd, input_data=strip_non_bmp(html).encode("utf-8"))


def fill_docx_template(docx_bytes: bytes, variables: dict) -> bytes:
    """Fill a DOCX template by replacing ``{{variable}}`` placeholders.

    Uses ``python-docx`` to iterate all paragraphs and table cells and
    perform find-replace on their text content.  Placeholders use the
    ``{{key}}`` syntax (double curly braces, no spaces).

    Parameters
    ----------
    docx_bytes:
        Raw DOCX bytes of the template file.
    variables:
        Mapping of placeholder names to replacement values.

    Returns
    -------
    bytes
        Modified DOCX bytes with placeholders replaced.

    Notes
    -----
    - Complex paragraph formatting (e.g. mixed bold/italic within a run
      that spans a placeholder boundary) may be partially flattened.
    - Jinja2 constructs (``{% if … %}``) are NOT supported in DOCX templates.
    """
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    _replace_in_docx(doc, variables)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes to PDF bytes using pandoc.

    Parameters
    ----------
    docx_bytes:
        Raw DOCX bytes.

    Returns
    -------
    bytes
        Raw PDF bytes.

    Raises
    ------
    FileNotFoundError
        If pandoc is not installed.
    RuntimeError
        If pandoc exits with a non-zero status.
    """
    cmd = [_PANDOC_BIN, "--from=docx", "--to=pdf", "--output=-"]
    return await _run_pandoc(cmd, input_data=docx_bytes)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _replace_in_docx(doc: "Document", variables: dict) -> None:  # type: ignore[name-defined]
    """Replace ``{{key}}`` placeholders in all paragraphs and table cells."""
    for para in doc.paragraphs:
        _replace_in_paragraph(para, variables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, variables)


def _replace_in_paragraph(para: "Paragraph", variables: dict) -> None:  # type: ignore[name-defined]
    """Replace placeholders in a single paragraph's runs.

    To handle placeholders that may be split across runs, we reconstruct
    the full paragraph text, perform replacements, then write back.
    """
    full_text = "".join(run.text for run in para.runs)
    if "{{" not in full_text:
        return

    new_text = full_text
    for key, value in variables.items():
        new_text = new_text.replace("{{" + key + "}}", str(value))

    if new_text == full_text:
        return

    # Write the entire new text into the first run, clear the rest
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""


async def _run_pandoc(cmd: list[str], input_data: bytes) -> bytes:
    """Run a pandoc subprocess asynchronously, feeding *input_data* via stdin.

    Returns stdout bytes on success.

    Raises
    ------
    FileNotFoundError
        If the pandoc binary is not found.
    RuntimeError
        If pandoc exits with a non-zero return code.
    asyncio.TimeoutError
        If the process exceeds ``_PANDOC_TIMEOUT`` seconds.
    """
    try:
        proc = await asyncio.create_subprocess_exec(
            *cmd,
            stdin=asyncio.subprocess.PIPE,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
    except FileNotFoundError:
        raise FileNotFoundError(
            f"pandoc binary not found. Ensure pandoc is installed (cmd: {cmd[0]!r})."
        )

    stdout, stderr = await asyncio.wait_for(
        proc.communicate(input=input_data),
        timeout=_PANDOC_TIMEOUT,
    )

    if proc.returncode != 0:
        err_msg = stderr.decode("utf-8", errors="replace").strip()
        raise RuntimeError(
            f"pandoc exited with code {proc.returncode}: {err_msg}"
        )

    log.debug("pandoc: cmd=%r produced %d bytes", cmd, len(stdout))
    return stdout


# ---------------------------------------------------------------------------
# CSV helpers
# ---------------------------------------------------------------------------

# Regex matching a Markdown pipe-table separator row (``|---|---|`` or ``---|---``).
_MD_TABLE_SEPARATOR_RE = re.compile(
    r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$"
)


def rows_to_csv(
    rows: list[dict],
    headers: Optional[list[str]] = None,
) -> bytes:
    """Render a list of dict rows as CSV bytes (``excel`` dialect, UTF-8 BOM).

    Parameters
    ----------
    rows:
        List of dictionaries — each one represents a row.
    headers:
        Optional explicit column order. If omitted, columns are derived from
        the union of keys across all rows (preserving first-seen order).

    Raises
    ------
    ValueError
        If *rows* is empty (and no *headers* supplied) or contains non-dict items.
    """
    import csv  # noqa: PLC0415

    if not isinstance(rows, list):
        raise ValueError("'rows' must be a list of dictionaries.")

    if headers is None:
        seen: dict[str, None] = {}
        for row in rows:
            if not isinstance(row, dict):
                raise ValueError("Each row must be a dictionary.")
            for k in row.keys():
                if k not in seen:
                    seen[str(k)] = None
        headers = list(seen.keys())

    if not headers:
        raise ValueError(
            "Cannot generate CSV: no columns inferred and no rows provided."
        )

    buf = io.StringIO()
    # UTF-8 BOM so Excel auto-detects the encoding.
    buf.write("\ufeff")
    writer = csv.DictWriter(buf, fieldnames=headers, dialect="excel", extrasaction="ignore")
    writer.writeheader()
    for row in rows:
        if not isinstance(row, dict):
            raise ValueError("Each row must be a dictionary.")
        writer.writerow({h: ("" if row.get(h) is None else str(row.get(h))) for h in headers})
    return buf.getvalue().encode("utf-8")


def markdown_table_to_csv(md_text: str) -> bytes:
    """Convert the first Markdown pipe-table found in *md_text* to CSV bytes.

    Recognises GitHub-flavoured pipe tables::

        | col1 | col2 |
        |------|------|
        | a    | b    |

    Returns CSV with header row preserved.

    Raises
    ------
    ValueError
        If no Markdown table is found in *md_text*.
    """
    lines = md_text.splitlines()
    for i, line in enumerate(lines):
        if i + 1 >= len(lines):
            break
        if "|" not in line:
            continue
        if not _MD_TABLE_SEPARATOR_RE.match(lines[i + 1]):
            continue
        # Found a header row at line *i*; collect rows until a blank line or
        # a line without ``|``.
        header_cells = _split_md_row(line)
        data_rows: list[list[str]] = []
        j = i + 2
        while j < len(lines):
            row_line = lines[j]
            if not row_line.strip() or "|" not in row_line:
                break
            data_rows.append(_split_md_row(row_line))
            j += 1

        rows = [
            {header_cells[k]: (cells[k] if k < len(cells) else "") for k in range(len(header_cells))}
            for cells in data_rows
        ]
        return rows_to_csv(rows, headers=header_cells)

    raise ValueError("No Markdown pipe-table found in the supplied content.")


def _split_md_row(line: str) -> list[str]:
    """Split a Markdown table row line into trimmed cell values."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


# ---------------------------------------------------------------------------
# ZIP / Pandoc-bundle helpers
# ---------------------------------------------------------------------------

# Hard limits on extracted ZIP bundles — protect against zip-bombs and
# resource exhaustion in the API container.
_ZIP_MAX_TOTAL_BYTES = 10 * 1024 * 1024  # 10 MB uncompressed
_ZIP_MAX_FILES = 50
_ZIP_MAX_PATH_LEN = 200

# Pandoc bundle conventions
_BUNDLE_DEFAULTS_FILENAMES = ("defaults.yaml", "defaults.yml")


def extract_template_zip(zip_bytes: bytes, dest_dir: str) -> str:
    """Safely extract a Pandoc bundle ZIP into *dest_dir*.

    Implements zip-slip protection: every entry's resolved absolute path
    must remain inside *dest_dir*. Enforces total-size, file-count and
    path-length limits to prevent zip-bombs.

    Parameters
    ----------
    zip_bytes:
        Raw bytes of the ZIP archive.
    dest_dir:
        Existing directory where files should be extracted.

    Returns
    -------
    str
        Absolute path to the directory that should be passed to pandoc as
        cwd / ``--resource-path``. If the archive contains a single
        top-level directory wrapping all entries, that subdirectory is
        returned; otherwise *dest_dir* itself is returned.

    Raises
    ------
    ValueError
        If the archive violates any safety limit or contains unsafe paths.
    """
    import os  # noqa: PLC0415
    import zipfile  # noqa: PLC0415

    dest_abs = os.path.realpath(dest_dir)
    if not os.path.isdir(dest_abs):
        raise ValueError(f"Destination directory does not exist: {dest_dir!r}")

    try:
        zf = zipfile.ZipFile(io.BytesIO(zip_bytes))
    except zipfile.BadZipFile as exc:
        raise ValueError(f"Invalid ZIP archive: {exc}") from exc

    with zf:
        infolist = zf.infolist()
        if len(infolist) > _ZIP_MAX_FILES:
            raise ValueError(
                f"ZIP bundle has too many entries ({len(infolist)} > {_ZIP_MAX_FILES})."
            )
        total_size = sum(info.file_size for info in infolist)
        if total_size > _ZIP_MAX_TOTAL_BYTES:
            raise ValueError(
                f"ZIP bundle uncompressed size {total_size} exceeds limit "
                f"{_ZIP_MAX_TOTAL_BYTES} bytes."
            )

        top_level: set[str] = set()
        for info in infolist:
            name = info.filename
            if not name or name.endswith("/"):
                # Directory entry — still validate path safety.
                pass
            if len(name) > _ZIP_MAX_PATH_LEN:
                raise ValueError(f"ZIP entry path too long: {name!r}")
            if name.startswith("/") or ".." in name.replace("\\", "/").split("/"):
                raise ValueError(f"Unsafe ZIP entry path: {name!r}")

            target = os.path.realpath(os.path.join(dest_abs, name))
            if not (target == dest_abs or target.startswith(dest_abs + os.sep)):
                raise ValueError(f"ZIP entry escapes destination: {name!r}")

            # Track the first path segment to detect a single-root wrapper dir.
            first_seg = name.replace("\\", "/").split("/", 1)[0]
            if first_seg:
                top_level.add(first_seg)

        zf.extractall(dest_abs)

    # If the archive wraps everything inside a single top-level directory,
    # return that directory so pandoc can find ``defaults.yaml`` directly.
    if len(top_level) == 1:
        only = next(iter(top_level))
        candidate = os.path.join(dest_abs, only)
        if os.path.isdir(candidate):
            return candidate
    return dest_abs


def find_bundle_defaults_file(bundle_dir: str) -> Optional[str]:
    """Locate ``defaults.yaml`` (or ``.yml``) inside a Pandoc bundle directory.

    Returns the absolute path or ``None`` if not found.
    """
    import os  # noqa: PLC0415

    for name in _BUNDLE_DEFAULTS_FILENAMES:
        candidate = os.path.join(bundle_dir, name)
        if os.path.isfile(candidate):
            return candidate
    return None


async def pandoc_run_with_defaults(
    input_md: str,
    bundle_dir: str,
    defaults_filename: str = "defaults.yaml",
) -> bytes:
    """Run pandoc with a ``--defaults=<defaults>`` file inside *bundle_dir*.

    The pandoc subprocess runs with ``cwd=bundle_dir`` so all relative paths
    (template, resource-path, includes, …) inside the defaults file resolve
    against the bundle.

    When the bundle contains a ``puppeteer-config.json`` (used by the
    ``diagram.lua`` Lua filter to call ``mmdc``), a *per-invocation* augmented
    Puppeteer config is written to a temporary directory that also contains a
    fresh Chromium user-data-dir.  This satisfies the chromium crashpad daemon
    requirement for a writable ``--database`` path and avoids the error
    ``chrome_crashpad_handler: --database is required`` that appears when the
    process runs with HOME unset or inside a read-only filesystem.

    Parameters
    ----------
    input_md:
        Markdown source to feed via stdin.
    bundle_dir:
        Absolute path to the bundle directory containing the defaults file.
    defaults_filename:
        Name of the defaults file inside the bundle. Defaults to ``defaults.yaml``.

    Returns
    -------
    bytes
        Raw bytes produced by pandoc on stdout (typically PDF).

    Raises
    ------
    FileNotFoundError
        If pandoc is not installed or the defaults file is missing.
    RuntimeError
        If pandoc exits with a non-zero status.
    """
    import json  # noqa: PLC0415
    import os  # noqa: PLC0415
    import shutil  # noqa: PLC0415
    import tempfile  # noqa: PLC0415

    defaults_path = os.path.join(bundle_dir, defaults_filename)
    if not os.path.isfile(defaults_path):
        raise FileNotFoundError(
            f"Pandoc defaults file not found in bundle: {defaults_path!r}"
        )

    cmd = [
        _PANDOC_BIN,
        f"--defaults={defaults_filename}",
        "--output=-",
    ]

    # Build environment for the pandoc subprocess. When the bundle includes
    # a ``puppeteer-config.json`` (used by the diagram.lua filter to launch
    # mmdc/headless Chromium with ``--no-sandbox`` etc.), create a per-call
    # augmented config that adds a writable ``--user-data-dir`` so
    # chromium's crashpad daemon finds a valid database path.
    env = os.environ.copy()
    static_config = os.path.join(bundle_dir, "puppeteer-config.json")

    chrome_home: str | None = None
    if os.path.isfile(static_config):
        chrome_home = tempfile.mkdtemp(prefix="gsage-pandoc-chrome-")
        user_data_dir = os.path.join(chrome_home, "profile")
        os.makedirs(user_data_dir, exist_ok=True)

        # Load the static args from the bundle config and append the
        # per-call user-data-dir so concurrent pandoc runs don't share a
        # Chromium profile (which would cause lock contention).
        try:
            with open(static_config, encoding="utf-8") as fh:
                static_args: list[str] = json.load(fh).get("args", [])
        except Exception:  # noqa: BLE001
            static_args = []

        dynamic_config_path = os.path.join(chrome_home, "puppeteer.json")
        with open(dynamic_config_path, "w", encoding="utf-8") as fh:
            json.dump(
                {
                    "args": [
                        *static_args,
                        f"--user-data-dir={user_data_dir}",
                    ]
                },
                fh,
            )

        env["PANDOC_DIAGRAM_PUPPETEER_CONFIG"] = dynamic_config_path
        # Chromium requires a writable HOME to place its crashpad socket.
        env["HOME"] = chrome_home
        env["XDG_CONFIG_HOME"] = chrome_home
        env["XDG_CACHE_HOME"] = chrome_home
        env["TMPDIR"] = chrome_home

    try:
        proc = await asyncio.create_subprocess_exec(
            *cmd,
            cwd=bundle_dir,
            stdin=asyncio.subprocess.PIPE,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
            env=env,
        )
    except FileNotFoundError:
        if chrome_home:
            shutil.rmtree(chrome_home, ignore_errors=True)
        raise FileNotFoundError(
            f"pandoc binary not found. Ensure pandoc is installed (cmd: {cmd[0]!r})."
        )

    try:
        stdout, stderr = await asyncio.wait_for(
            proc.communicate(input=strip_non_bmp(input_md).encode("utf-8")),
            timeout=_PANDOC_TIMEOUT * 2,  # bundle/LaTeX runs are slower
        )
    finally:
        if chrome_home:
            shutil.rmtree(chrome_home, ignore_errors=True)

    if proc.returncode != 0:
        err_msg = stderr.decode("utf-8", errors="replace").strip()
        raise RuntimeError(
            f"pandoc exited with code {proc.returncode}: {err_msg}"
        )

    log.debug(
        "pandoc-bundle: cwd=%s defaults=%s produced %d bytes",
        bundle_dir, defaults_filename, len(stdout),
    )
    return stdout
